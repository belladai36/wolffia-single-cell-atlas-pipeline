#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent
INPUT_PATH = PROJECT_ROOT / "docs" / "wolffia_data_generation_protocol.md"
OUTPUT_PATH = PROJECT_ROOT / "output" / "docx" / "wolffia_data_generation_protocol.docx"


def set_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_numbered_paragraph(paragraph, level: int = 0) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "9")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def build_docx(lines: list[str]) -> Document:
    doc = Document()
    set_margins(doc)
    configure_styles(doc)

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        line = raw.strip()
        if not line:
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_inline(line[2:]))
            run.bold = True
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.size = Pt(18)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_paragraph(clean_inline(line[3:]), style="Heading 1")
            i += 1
            continue

        if line.startswith("### "):
            doc.add_paragraph(clean_inline(line[4:]), style="Heading 2")
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(clean_inline(line.lstrip("> ").strip()))
            run.italic = True
            i += 1
            continue

        if line.startswith("- "):
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt.startswith("- "):
                    break
                p = doc.add_paragraph(style="Normal")
                p.style = doc.styles["List Bullet"]
                p.add_run(clean_inline(nxt[2:].strip()))
                i += 1
            continue

        if len(line) > 2 and line[0].isdigit() and line[1] == ".":
            while i < len(lines):
                nxt = lines[i].strip()
                if not (len(nxt) > 2 and nxt[0].isdigit() and nxt[1] == "."):
                    break
                p = doc.add_paragraph(style="List Number")
                p.add_run(clean_inline(nxt[2:].strip()))
                i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith(("# ", "## ", "### ", "- ", ">")):
                break
            if len(nxt) > 2 and nxt[0].isdigit() and nxt[1] == ".":
                break
            paragraph_lines.append(nxt)
            i += 1

        p = doc.add_paragraph(clean_inline(" ".join(paragraph_lines)))
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15

    return doc


def main():
    lines = INPUT_PATH.read_text(encoding="utf-8").splitlines()
    document = build_docx(lines)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
